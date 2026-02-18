"""
Criminal Appeal Case Management System
A comprehensive Flask application for managing criminal appeal cases with document management,
timeline tracking, legal analysis, and barrister report generation.
"""

import os
import logging
import mimetypes
from datetime import datetime
from werkzeug.utils import secure_filename
from pathlib import Path

from flask import Flask, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_marshmallow import Marshmallow
from sqlalchemy import or_, desc
from sqlalchemy.exc import SQLAlchemyError, IntegrityError

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

# Initialize Flask app
app = Flask(__name__)

# Configure logging early
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
)
logger = logging.getLogger(__name__)

# Configuration
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', 'sqlite:///criminal_appeal.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size
app.config['UPLOAD_FOLDER'] = os.path.join(os.path.dirname(__file__), 'uploads')
app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'docx', 'txt'}

# SECRET_KEY is required for production - fail if not set
secret_key = os.environ.get('SECRET_KEY')
if not secret_key:
    if os.environ.get('FLASK_ENV') == 'development':
        logger.warning('Using default SECRET_KEY for development. DO NOT use in production!')
        secret_key = 'dev-secret-key-DO-NOT-USE-IN-PRODUCTION'
    else:
        raise ValueError('SECRET_KEY environment variable must be set in production')
app.config['SECRET_KEY'] = secret_key

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Initialize extensions
db = SQLAlchemy(app)
ma = Marshmallow(app)


# =====================================================================
# DATABASE MODELS
# =====================================================================

class Case(db.Model):
    """
    Case model representing a criminal appeal case.
    """
    __tablename__ = 'cases'
    
    id = db.Column(db.Integer, primary_key=True)
    case_number = db.Column(db.String(100), unique=True, nullable=False, index=True)
    defendant_name = db.Column(db.String(200), nullable=False)
    offense_type = db.Column(db.String(200), nullable=False, default='Murder')
    court = db.Column(db.String(200), nullable=False, default='NSW Supreme Court')
    status = db.Column(db.String(50), nullable=False, default='Open')
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow, onupdate=datetime.utcnow)
    
    # Relationships
    documents = db.relationship('Document', backref='case', lazy=True, cascade='all, delete-orphan')
    timeline_events = db.relationship('TimelineEvent', backref='case', lazy=True, cascade='all, delete-orphan')
    legal_analyses = db.relationship('LegalAnalysis', backref='case', lazy=True, cascade='all, delete-orphan')
    barrister_reports = db.relationship('BarristerReport', backref='case', lazy=True, cascade='all, delete-orphan')

    def __repr__(self):
        return f'<Case {self.case_number}>'


class Document(db.Model):
    """
    Document model representing uploaded files associated with a case.
    """
    __tablename__ = 'documents'
    
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    title = db.Column(db.String(255), nullable=False)
    document_type = db.Column(db.String(100), nullable=False)
    file_path = db.Column(db.String(500), nullable=False)
    file_type = db.Column(db.String(10), nullable=False)
    file_size = db.Column(db.Integer, nullable=False)
    extracted_text = db.Column(db.Text, nullable=True)
    upload_date = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    
    # Relationships
    timeline_events = db.relationship('TimelineEvent', backref='document', lazy=True, cascade='all, delete-orphan')

    def __repr__(self):
        return f'<Document {self.title}>'


class TimelineEvent(db.Model):
    """
    TimelineEvent model representing significant events in a case timeline.
    """
    __tablename__ = 'timeline_events'
    
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    document_id = db.Column(db.Integer, db.ForeignKey('documents.id', ondelete='SET NULL'), nullable=True)
    event_date = db.Column(db.DateTime, nullable=False)
    event_type = db.Column(db.String(100), nullable=False)
    description = db.Column(db.Text, nullable=False)
    significance = db.Column(db.String(50), nullable=False, default='Medium')
    relevance_to_appeal = db.Column(db.Text, nullable=True)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)

    def __repr__(self):
        return f'<TimelineEvent {self.event_type} - {self.event_date}>'


class LegalAnalysis(db.Model):
    """
    LegalAnalysis model representing legal analysis of grounds for appeal.
    """
    __tablename__ = 'legal_analyses'
    
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    ground_of_merit = db.Column(db.String(255), nullable=False)
    legal_basis = db.Column(db.Text, nullable=False)
    strength_assessment = db.Column(db.String(50), nullable=False)
    nsw_law_references = db.Column(db.Text, nullable=True)
    federal_law_references = db.Column(db.Text, nullable=True)
    supporting_evidence = db.Column(db.Text, nullable=True)
    analysis_summary = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow, onupdate=datetime.utcnow)

    def __repr__(self):
        return f'<LegalAnalysis {self.ground_of_merit}>'


class BarristerReport(db.Model):
    """
    BarristerReport model representing generated barrister reports.
    """
    __tablename__ = 'barrister_reports'
    
    id = db.Column(db.Integer, primary_key=True)
    case_id = db.Column(db.Integer, db.ForeignKey('cases.id', ondelete='CASCADE'), nullable=False)
    report_title = db.Column(db.String(255), nullable=False)
    executive_summary = db.Column(db.Text, nullable=False)
    grounds_identified = db.Column(db.Text, nullable=False)
    legal_analysis_summary = db.Column(db.Text, nullable=False)
    recommendations = db.Column(db.Text, nullable=False)
    generated_at = db.Column(db.DateTime, nullable=False, default=datetime.utcnow)
    generated_by = db.Column(db.String(200), nullable=False)

    def __repr__(self):
        return f'<BarristerReport {self.report_title}>'


# =====================================================================
# MARSHMALLOW SCHEMAS
# =====================================================================

class CaseSchema(ma.SQLAlchemyAutoSchema):
    """Schema for Case model serialization."""
    class Meta:
        model = Case
        include_fk = True
        load_instance = True
        exclude = ('documents', 'timeline_events', 'legal_analyses', 'barrister_reports')


class DocumentSchema(ma.SQLAlchemyAutoSchema):
    """Schema for Document model serialization."""
    class Meta:
        model = Document
        include_fk = True
        load_instance = True
        exclude = ('case',)


class TimelineEventSchema(ma.SQLAlchemyAutoSchema):
    """Schema for TimelineEvent model serialization."""
    class Meta:
        model = TimelineEvent
        include_fk = True
        load_instance = True
        exclude = ('case', 'document')


class LegalAnalysisSchema(ma.SQLAlchemyAutoSchema):
    """Schema for LegalAnalysis model serialization."""
    class Meta:
        model = LegalAnalysis
        include_fk = True
        load_instance = True
        exclude = ('case',)


class BarristerReportSchema(ma.SQLAlchemyAutoSchema):
    """Schema for BarristerReport model serialization."""
    class Meta:
        model = BarristerReport
        include_fk = True
        load_instance = True
        exclude = ('case',)


# Initialize schemas
case_schema = CaseSchema()
cases_schema = CaseSchema(many=True)
document_schema = DocumentSchema()
documents_schema = DocumentSchema(many=True)
timeline_event_schema = TimelineEventSchema()
timeline_events_schema = TimelineEventSchema(many=True)
legal_analysis_schema = LegalAnalysisSchema()
legal_analyses_schema = LegalAnalysisSchema(many=True)
barrister_report_schema = BarristerReportSchema()
barrister_reports_schema = BarristerReportSchema(many=True)


# =====================================================================
# HELPER FUNCTIONS
# =====================================================================

def parse_datetime(date_string):
    """
    Robustly parse datetime string in various ISO 8601 formats.
    
    Args:
        date_string (str): Date string to parse
        
    Returns:
        datetime: Parsed datetime object
        
    Raises:
        ValueError: If the date string cannot be parsed
    """
    if not date_string:
        raise ValueError("Date string cannot be empty")
    
    # Handle common ISO 8601 variations
    # Replace 'Z' with '+00:00' for timezone-aware parsing
    if date_string.endswith('Z'):
        date_string = date_string[:-1] + '+00:00'
    
    try:
        return datetime.fromisoformat(date_string)
    except ValueError as e:
        # Try alternative formats
        for fmt in ['%Y-%m-%dT%H:%M:%S', '%Y-%m-%d %H:%M:%S', '%Y-%m-%d']:
            try:
                return datetime.strptime(date_string, fmt)
            except ValueError:
                continue
        
        raise ValueError(f"Unable to parse date string: {date_string}. Use ISO 8601 format (e.g., '2024-01-01T12:00:00' or '2024-01-01T12:00:00Z')")


def allowed_file(filename):
    """
    Check if the file extension is allowed.
    
    Args:
        filename (str): The filename to check
        
    Returns:
        bool: True if file extension is allowed, False otherwise
    """
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def extract_text_from_pdf(file_path):
    """
    Extract text content from a PDF file using PyPDF2.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content or error message
    """
    if PyPDF2 is None:
        logger.warning("PyPDF2 not installed, cannot extract PDF text")
        return "PyPDF2 not installed - text extraction unavailable"
    
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        
        logger.info(f"Successfully extracted text from PDF: {file_path}")
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return f"Error extracting text: {str(e)}"


def extract_text_from_docx(file_path):
    """
    Extract text content from a DOCX file using python-docx.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text content or error message
    """
    if docx is None:
        logger.warning("python-docx not installed, cannot extract DOCX text")
        return "python-docx not installed - text extraction unavailable"
    
    try:
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        
        logger.info(f"Successfully extracted text from DOCX: {file_path}")
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return f"Error extracting text: {str(e)}"


def extract_text_from_txt(file_path):
    """
    Extract text content from a TXT file.
    
    Args:
        file_path (str): Path to the TXT file
        
    Returns:
        str: Extracted text content or error message
    """
    try:
        # Try UTF-8 first
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Fallback to other common encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.info(f"Successfully read file with {encoding} encoding: {file_path}")
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise UnicodeDecodeError('multiple', b'', 0, 0, 'Failed to decode file with common encodings')
        
        logger.info(f"Successfully extracted text from TXT: {file_path}")
        return text.strip()
    except UnicodeDecodeError as e:
        logger.error(f"Unicode decoding error for TXT {file_path}: {str(e)}")
        return f"Error: Unable to decode file. File may not be a text file or uses an unsupported encoding."
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
        return f"Error extracting text: {str(e)}"


def extract_text_from_file(file_path, file_type):
    """
    Extract text from a file based on its type.
    
    Args:
        file_path (str): Path to the file
        file_type (str): Type of file (pdf, docx, txt)
        
    Returns:
        str: Extracted text content
    """
    if file_type == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'txt':
        return extract_text_from_txt(file_path)
    else:
        return "Unsupported file type for text extraction"


def auto_create_timeline_event(document):
    """
    Automatically create a timeline event when a document is uploaded.
    
    Args:
        document (Document): The uploaded document object
        
    Returns:
        TimelineEvent: The created timeline event
    """
    try:
        event = TimelineEvent(
            case_id=document.case_id,
            document_id=document.id,
            event_date=document.upload_date,
            event_type='Document Upload',
            description=f"Document '{document.title}' ({document.document_type}) was uploaded to the case",
            significance='Medium',
            relevance_to_appeal=f"This {document.document_type} may contain relevant information for the appeal"
        )
        
        db.session.add(event)
        db.session.commit()
        
        logger.info(f"Auto-created timeline event for document {document.id}")
        return event
    except Exception as e:
        logger.error(f"Error creating timeline event: {str(e)}")
        db.session.rollback()
        return None


def analyze_grounds_of_merit(case):
    """
    Analyze potential grounds of merit for an appeal.
    
    **PLACEHOLDER IMPLEMENTATION**: This function currently returns sample analyses.
    In production, this should be replaced with actual AI-powered analysis using OpenAI API
    or similar NLP service to analyze case documents and identify real grounds of merit.
    
    TODO: Integrate with OpenAI API to analyze extracted text from case documents
    TODO: Implement NLP-based extraction of relevant legal issues
    TODO: Implement case law matching and precedent analysis
    
    Args:
        case (Case): The case to analyze
        
    Returns:
        list: List of created LegalAnalysis objects
    """
    analyses = []
    
    # PLACEHOLDER: Sample grounds of merit with NSW law references
    # In production, these would be generated by analyzing actual case documents
    sample_grounds = [
        {
            'ground': 'Error in Law - Judicial Misdirection',
            'legal_basis': 'The trial judge misdirected the jury on a material element of the offense, leading to a miscarriage of justice.',
            'strength': 'Strong',
            'nsw_refs': 'Criminal Appeal Act 1912 (NSW) s 6(1); Crimes (Appeal and Review) Act 2001 (NSW) s 53',
            'federal_refs': 'None applicable',
            'evidence': 'Trial transcript pages 45-48 showing erroneous jury directions',
            'summary': 'Strong ground based on clear judicial error in directing the jury on the element of intent. NSW Court of Criminal Appeal has consistently held that such misdirections constitute a material irregularity.'
        },
        {
            'ground': 'Fresh Evidence',
            'legal_basis': 'New evidence has come to light that was not available at trial and could reasonably affect the outcome.',
            'strength': 'Medium',
            'nsw_refs': 'Criminal Appeal Act 1912 (NSW) s 12; Crimes (Appeal and Review) Act 2001 (NSW) s 54',
            'federal_refs': 'Evidence Act 1995 (Cth) s 97, s 101',
            'evidence': 'Witness statement dated after trial conclusion; forensic analysis report',
            'summary': 'Medium strength ground. The fresh evidence meets the threshold of credibility but requires demonstration that it could not have been discovered with reasonable diligence at trial.'
        },
        {
            'ground': 'Unreasonable Verdict',
            'legal_basis': 'The verdict is unreasonable or cannot be supported having regard to the evidence presented at trial.',
            'strength': 'Medium',
            'nsw_refs': 'Criminal Appeal Act 1912 (NSW) s 6(1); Crimes (Appeal and Review) Act 2001 (NSW) s 53(1)(a)',
            'federal_refs': 'None applicable',
            'evidence': 'Contradictory witness testimony; lack of corroborating evidence',
            'summary': 'Medium strength ground. Requires detailed analysis of the trial record to demonstrate that no reasonable jury, properly instructed, could have reached the guilty verdict on the evidence presented.'
        }
    ]
    
    try:
        for ground_data in sample_grounds:
            analysis = LegalAnalysis(
                case_id=case.id,
                ground_of_merit=ground_data['ground'],
                legal_basis=ground_data['legal_basis'],
                strength_assessment=ground_data['strength'],
                nsw_law_references=ground_data['nsw_refs'],
                federal_law_references=ground_data['federal_refs'],
                supporting_evidence=ground_data['evidence'],
                analysis_summary=ground_data['summary']
            )
            
            db.session.add(analysis)
            analyses.append(analysis)
        
        db.session.commit()
        logger.info(f"Created {len(analyses)} legal analyses for case {case.id}")
        return analyses
    except Exception as e:
        logger.error(f"Error analyzing grounds of merit: {str(e)}")
        db.session.rollback()
        return []


# =====================================================================
# API ENDPOINTS - CASES
# =====================================================================

@app.route('/api/cases', methods=['POST'])
def create_case():
    """
    Create a new case.
    
    Request Body:
        {
            "case_number": str,
            "defendant_name": str,
            "offense_type": str,
            "court": str,
            "status": str (optional)
        }
    
    Returns:
        JSON response with created case data
    """
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No input data provided'}), 400
        
        # Validate required fields
        required_fields = ['case_number', 'defendant_name']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        # Check if case number already exists
        existing_case = Case.query.filter_by(case_number=data['case_number']).first()
        if existing_case:
            return jsonify({'error': 'Case number already exists'}), 409
        
        new_case = Case(
            case_number=data['case_number'],
            defendant_name=data['defendant_name'],
            offense_type=data.get('offense_type', 'Murder'),
            court=data.get('court', 'NSW Supreme Court'),
            status=data.get('status', 'Open')
        )
        
        db.session.add(new_case)
        db.session.commit()
        
        logger.info(f"Created new case: {new_case.case_number}")
        return jsonify(case_schema.dump(new_case)), 201
        
    except IntegrityError as e:
        db.session.rollback()
        logger.error(f"Integrity error creating case: {str(e)}")
        return jsonify({'error': 'Database integrity error', 'details': str(e)}), 400
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error creating case: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating case: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases', methods=['GET'])
def get_cases():
    """
    Get all cases with optional filtering.
    
    Query Parameters:
        status: Filter by case status
        search: Search in case_number, defendant_name, or offense_type
        
    Returns:
        JSON array of cases
    """
    try:
        query = Case.query
        
        # Filter by status
        status = request.args.get('status')
        if status:
            query = query.filter_by(status=status)
        
        # Search functionality
        search = request.args.get('search')
        if search:
            search_filter = or_(
                Case.case_number.ilike(f'%{search}%'),
                Case.defendant_name.ilike(f'%{search}%'),
                Case.offense_type.ilike(f'%{search}%')
            )
            query = query.filter(search_filter)
        
        cases = query.order_by(desc(Case.created_at)).all()
        
        logger.info(f"Retrieved {len(cases)} cases")
        return jsonify(cases_schema.dump(cases)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving cases: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving cases: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>', methods=['GET'])
def get_case(case_id):
    """
    Get a specific case by ID.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON response with case data
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        logger.info(f"Retrieved case: {case.case_number}")
        return jsonify(case_schema.dump(case)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving case {case_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving case {case_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>', methods=['PUT'])
def update_case(case_id):
    """
    Update a case.
    
    Args:
        case_id: The ID of the case to update
        
    Returns:
        JSON response with updated case data
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No input data provided'}), 400
        
        # Update allowed fields
        if 'defendant_name' in data:
            case.defendant_name = data['defendant_name']
        if 'offense_type' in data:
            case.offense_type = data['offense_type']
        if 'court' in data:
            case.court = data['court']
        if 'status' in data:
            case.status = data['status']
        
        case.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        logger.info(f"Updated case: {case.case_number}")
        return jsonify(case_schema.dump(case)), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error updating case {case_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating case {case_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>', methods=['DELETE'])
def delete_case(case_id):
    """
    Delete a case and all associated data.
    
    Args:
        case_id: The ID of the case to delete
        
    Returns:
        JSON response confirming deletion
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        case_number = case.case_number
        
        # Delete associated files
        for document in case.documents:
            if os.path.exists(document.file_path):
                try:
                    os.remove(document.file_path)
                except OSError as e:
                    logger.warning(f"Could not delete file {document.file_path}: {str(e)}")
        
        db.session.delete(case)
        db.session.commit()
        
        logger.info(f"Deleted case: {case_number}")
        return jsonify({'message': f'Case {case_number} deleted successfully'}), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error deleting case {case_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting case {case_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


# =====================================================================
# API ENDPOINTS - DOCUMENTS
# =====================================================================

@app.route('/api/cases/<int:case_id>/documents', methods=['POST'])
def upload_document(case_id):
    """
    Upload a document to a case.
    
    Args:
        case_id: The ID of the case
        
    Form Data:
        file: The file to upload
        title: Document title
        document_type: Type of document
        
    Returns:
        JSON response with created document data
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'error': 'File type not allowed. Allowed types: PDF, DOCX, TXT'}), 400
        
        title = request.form.get('title')
        document_type = request.form.get('document_type', 'Other')
        
        if not title:
            return jsonify({'error': 'Document title is required'}), 400
        
        # Secure the filename and validate
        filename = secure_filename(file.filename)
        if not filename or '.' not in filename:
            return jsonify({'error': 'Invalid filename. Please provide a valid file with an extension.'}), 400
        
        file_extension = filename.rsplit('.', 1)[1].lower()
        
        # Create case-specific directory
        case_dir = os.path.join(app.config['UPLOAD_FOLDER'], f"case_{case_id}")
        os.makedirs(case_dir, exist_ok=True)
        
        # Generate unique filename
        timestamp = datetime.utcnow().strftime('%Y%m%d_%H%M%S')
        unique_filename = f"{timestamp}_{filename}"
        file_path = os.path.join(case_dir, unique_filename)
        
        # Save file
        file.save(file_path)
        file_size = os.path.getsize(file_path)
        
        # Extract text
        extracted_text = extract_text_from_file(file_path, file_extension)
        
        # Create document record
        document = Document(
            case_id=case_id,
            title=title,
            document_type=document_type,
            file_path=file_path,
            file_type=file_extension,
            file_size=file_size,
            extracted_text=extracted_text
        )
        
        db.session.add(document)
        db.session.commit()
        
        # Auto-create timeline event
        auto_create_timeline_event(document)
        
        logger.info(f"Uploaded document {document.id} to case {case_id}")
        return jsonify(document_schema.dump(document)), 201
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error uploading document: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error uploading document: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>/documents', methods=['GET'])
def get_case_documents(case_id):
    """
    Get all documents for a case.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON array of documents
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        documents = Document.query.filter_by(case_id=case_id).order_by(desc(Document.upload_date)).all()
        
        logger.info(f"Retrieved {len(documents)} documents for case {case_id}")
        return jsonify(documents_schema.dump(documents)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving documents: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving documents: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/documents/<int:document_id>', methods=['GET'])
def get_document(document_id):
    """
    Get a specific document by ID.
    
    Args:
        document_id: The ID of the document
        
    Returns:
        JSON response with document data
    """
    try:
        document = Document.query.get(document_id)
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        logger.info(f"Retrieved document {document_id}")
        return jsonify(document_schema.dump(document)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving document {document_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving document {document_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/documents/<int:document_id>/download', methods=['GET'])
def download_document(document_id):
    """
    Download a document file.
    
    Args:
        document_id: The ID of the document
        
    Returns:
        File download
    """
    try:
        document = Document.query.get(document_id)
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        if not os.path.exists(document.file_path):
            return jsonify({'error': 'File not found on server'}), 404
        
        logger.info(f"Downloading document {document_id}")
        return send_file(
            document.file_path,
            as_attachment=True,
            download_name=f"{document.title}.{document.file_type}"
        )
        
    except Exception as e:
        logger.error(f"Error downloading document {document_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/documents/<int:document_id>', methods=['DELETE'])
def delete_document(document_id):
    """
    Delete a document.
    
    Args:
        document_id: The ID of the document to delete
        
    Returns:
        JSON response confirming deletion
    """
    try:
        document = Document.query.get(document_id)
        
        if not document:
            return jsonify({'error': 'Document not found'}), 404
        
        file_path = document.file_path
        
        # Try to delete file first - if it fails, don't delete database record
        file_deleted = False
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                file_deleted = True
                logger.info(f"Deleted file: {file_path}")
            except OSError as e:
                logger.error(f"Failed to delete file {file_path}: {str(e)}")
                return jsonify({'error': 'Failed to delete file from disk', 'details': str(e)}), 500
        else:
            # File doesn't exist, but we can still delete the DB record
            logger.warning(f"File not found on disk: {file_path}")
            file_deleted = True
        
        # Only delete database record if file was successfully deleted or didn't exist
        if file_deleted:
            db.session.delete(document)
            db.session.commit()
            
            logger.info(f"Deleted document {document_id}")
            return jsonify({'message': 'Document deleted successfully'}), 200
        else:
            return jsonify({'error': 'Failed to delete document'}), 500
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error deleting document {document_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting document {document_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


# =====================================================================
# API ENDPOINTS - TIMELINE
# =====================================================================

@app.route('/api/cases/<int:case_id>/timeline', methods=['POST'])
def create_timeline_event(case_id):
    """
    Create a timeline event for a case.
    
    Args:
        case_id: The ID of the case
        
    Request Body:
        {
            "event_date": ISO date string,
            "event_type": str,
            "description": str,
            "significance": str (optional),
            "relevance_to_appeal": str (optional),
            "document_id": int (optional)
        }
        
    Returns:
        JSON response with created timeline event
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No input data provided'}), 400
        
        required_fields = ['event_date', 'event_type', 'description']
        for field in required_fields:
            if field not in data:
                return jsonify({'error': f'Missing required field: {field}'}), 400
        
        # Parse event date
        try:
            event_date = parse_datetime(data['event_date'])
        except ValueError as e:
            return jsonify({'error': str(e)}), 400
        
        event = TimelineEvent(
            case_id=case_id,
            document_id=data.get('document_id'),
            event_date=event_date,
            event_type=data['event_type'],
            description=data['description'],
            significance=data.get('significance', 'Medium'),
            relevance_to_appeal=data.get('relevance_to_appeal')
        )
        
        db.session.add(event)
        db.session.commit()
        
        logger.info(f"Created timeline event {event.id} for case {case_id}")
        return jsonify(timeline_event_schema.dump(event)), 201
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error creating timeline event: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating timeline event: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>/timeline', methods=['GET'])
def get_case_timeline(case_id):
    """
    Get all timeline events for a case.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON array of timeline events
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        events = TimelineEvent.query.filter_by(case_id=case_id).order_by(TimelineEvent.event_date).all()
        
        logger.info(f"Retrieved {len(events)} timeline events for case {case_id}")
        return jsonify(timeline_events_schema.dump(events)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving timeline events: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving timeline events: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/timeline/<int:event_id>', methods=['GET'])
def get_timeline_event(event_id):
    """
    Get a specific timeline event by ID.
    
    Args:
        event_id: The ID of the timeline event
        
    Returns:
        JSON response with timeline event data
    """
    try:
        event = TimelineEvent.query.get(event_id)
        
        if not event:
            return jsonify({'error': 'Timeline event not found'}), 404
        
        logger.info(f"Retrieved timeline event {event_id}")
        return jsonify(timeline_event_schema.dump(event)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/timeline/<int:event_id>', methods=['PUT'])
def update_timeline_event(event_id):
    """
    Update a timeline event.
    
    Args:
        event_id: The ID of the timeline event to update
        
    Returns:
        JSON response with updated timeline event data
    """
    try:
        event = TimelineEvent.query.get(event_id)
        
        if not event:
            return jsonify({'error': 'Timeline event not found'}), 404
        
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No input data provided'}), 400
        
        if 'event_date' in data:
            try:
                event.event_date = parse_datetime(data['event_date'])
            except ValueError as e:
                return jsonify({'error': str(e)}), 400
                return jsonify({'error': 'Invalid date format'}), 400
        
        if 'event_type' in data:
            event.event_type = data['event_type']
        if 'description' in data:
            event.description = data['description']
        if 'significance' in data:
            event.significance = data['significance']
        if 'relevance_to_appeal' in data:
            event.relevance_to_appeal = data['relevance_to_appeal']
        
        db.session.commit()
        
        logger.info(f"Updated timeline event {event_id}")
        return jsonify(timeline_event_schema.dump(event)), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error updating timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/timeline/<int:event_id>', methods=['DELETE'])
def delete_timeline_event(event_id):
    """
    Delete a timeline event.
    
    Args:
        event_id: The ID of the timeline event to delete
        
    Returns:
        JSON response confirming deletion
    """
    try:
        event = TimelineEvent.query.get(event_id)
        
        if not event:
            return jsonify({'error': 'Timeline event not found'}), 404
        
        db.session.delete(event)
        db.session.commit()
        
        logger.info(f"Deleted timeline event {event_id}")
        return jsonify({'message': 'Timeline event deleted successfully'}), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error deleting timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting timeline event {event_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


# =====================================================================
# API ENDPOINTS - LEGAL ANALYSIS
# =====================================================================

@app.route('/api/cases/<int:case_id>/analyze', methods=['POST'])
def analyze_case(case_id):
    """
    Analyze a case for potential grounds of merit.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON array of created legal analyses
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        analyses = analyze_grounds_of_merit(case)
        
        if not analyses:
            return jsonify({'error': 'Failed to create analyses'}), 500
        
        logger.info(f"Analyzed case {case_id}, created {len(analyses)} legal analyses")
        return jsonify(legal_analyses_schema.dump(analyses)), 201
        
    except Exception as e:
        logger.error(f"Error analyzing case {case_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>/analyses', methods=['GET'])
def get_case_analyses(case_id):
    """
    Get all legal analyses for a case.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON array of legal analyses
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        analyses = LegalAnalysis.query.filter_by(case_id=case_id).order_by(desc(LegalAnalysis.created_at)).all()
        
        logger.info(f"Retrieved {len(analyses)} legal analyses for case {case_id}")
        return jsonify(legal_analyses_schema.dump(analyses)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving legal analyses: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving legal analyses: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/analyses/<int:analysis_id>', methods=['GET'])
def get_analysis(analysis_id):
    """
    Get a specific legal analysis by ID.
    
    Args:
        analysis_id: The ID of the legal analysis
        
    Returns:
        JSON response with legal analysis data
    """
    try:
        analysis = LegalAnalysis.query.get(analysis_id)
        
        if not analysis:
            return jsonify({'error': 'Legal analysis not found'}), 404
        
        logger.info(f"Retrieved legal analysis {analysis_id}")
        return jsonify(legal_analysis_schema.dump(analysis)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/analyses/<int:analysis_id>', methods=['PUT'])
def update_analysis(analysis_id):
    """
    Update a legal analysis.
    
    Args:
        analysis_id: The ID of the legal analysis to update
        
    Returns:
        JSON response with updated legal analysis data
    """
    try:
        analysis = LegalAnalysis.query.get(analysis_id)
        
        if not analysis:
            return jsonify({'error': 'Legal analysis not found'}), 404
        
        data = request.get_json()
        
        if not data:
            return jsonify({'error': 'No input data provided'}), 400
        
        if 'ground_of_merit' in data:
            analysis.ground_of_merit = data['ground_of_merit']
        if 'legal_basis' in data:
            analysis.legal_basis = data['legal_basis']
        if 'strength_assessment' in data:
            analysis.strength_assessment = data['strength_assessment']
        if 'nsw_law_references' in data:
            analysis.nsw_law_references = data['nsw_law_references']
        if 'federal_law_references' in data:
            analysis.federal_law_references = data['federal_law_references']
        if 'supporting_evidence' in data:
            analysis.supporting_evidence = data['supporting_evidence']
        if 'analysis_summary' in data:
            analysis.analysis_summary = data['analysis_summary']
        
        analysis.updated_at = datetime.utcnow()
        
        db.session.commit()
        
        logger.info(f"Updated legal analysis {analysis_id}")
        return jsonify(legal_analysis_schema.dump(analysis)), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error updating legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error updating legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/analyses/<int:analysis_id>', methods=['DELETE'])
def delete_analysis(analysis_id):
    """
    Delete a legal analysis.
    
    Args:
        analysis_id: The ID of the legal analysis to delete
        
    Returns:
        JSON response confirming deletion
    """
    try:
        analysis = LegalAnalysis.query.get(analysis_id)
        
        if not analysis:
            return jsonify({'error': 'Legal analysis not found'}), 404
        
        db.session.delete(analysis)
        db.session.commit()
        
        logger.info(f"Deleted legal analysis {analysis_id}")
        return jsonify({'message': 'Legal analysis deleted successfully'}), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error deleting legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting legal analysis {analysis_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


# =====================================================================
# API ENDPOINTS - BARRISTER REPORTS
# =====================================================================

@app.route('/api/cases/<int:case_id>/reports', methods=['POST'])
def create_barrister_report(case_id):
    """
    Generate a professional barrister report for a case based on legal analyses.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON response with generated barrister report
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        # Get all legal analyses for the case
        analyses = LegalAnalysis.query.filter_by(case_id=case_id).all()
        
        if not analyses:
            return jsonify({'error': 'No legal analyses found. Please run analysis first.'}), 400
        
        # Auto-generate report content
        report_title = f"Legal Opinion - Appeal Prospects: {case.defendant_name} - Case {case.case_number}"
        
        executive_summary = f"""This report provides a comprehensive legal analysis of the appeal prospects for {case.defendant_name} 
in relation to the {case.offense_type} conviction. The analysis is based on review of all available case documents, 
evidence, and applicable NSW and Federal legislation. The analysis was conducted on {datetime.utcnow().strftime('%B %d, %Y')}."""
        
        # Build grounds identified section
        grounds_list = []
        for idx, analysis in enumerate(analyses, 1):
            grounds_list.append(f"{idx}. {analysis.ground_of_merit} (Strength: {analysis.strength_assessment})")
        grounds_identified = "\n".join(grounds_list)
        
        # Build legal analysis summary
        analysis_summary_parts = []
        for idx, analysis in enumerate(analyses, 1):
            summary = f"""
Ground {idx}: {analysis.ground_of_merit}

Legal Basis: {analysis.legal_basis}

Strength Assessment: {analysis.strength_assessment}

NSW Law References: {analysis.nsw_law_references or 'N/A'}

Federal Law References: {analysis.federal_law_references or 'N/A'}

Supporting Evidence: {analysis.supporting_evidence or 'N/A'}
"""
            analysis_summary_parts.append(summary)
        
        legal_analysis_summary = "\n---\n".join(analysis_summary_parts)
        
        # Generate recommendations based on strength of grounds
        strong_grounds = [a for a in analyses if a.strength_assessment == 'Strong']
        medium_grounds = [a for a in analyses if a.strength_assessment == 'Medium']
        
        if strong_grounds:
            recommendations = f"""Based on the analysis, there are {len(strong_grounds)} strong ground(s) of appeal identified. 
It is recommended to proceed with the appeal, focusing primarily on these strong grounds. 
The case presents reasonable prospects of success. Immediate steps should include:

1. Prepare detailed written submissions on the strong grounds
2. Compile all supporting documentary evidence
3. Engage expert witnesses where necessary
4. Consider application for leave to appeal if required

The {len(strong_grounds)} strong ground(s) provide a solid foundation for appellate review."""
        elif medium_grounds:
            recommendations = f"""The analysis has identified {len(medium_grounds)} medium strength ground(s) of appeal. 
While these grounds have merit, further investigation and evidence gathering is recommended 
to strengthen the appeal prospects before proceeding. Recommended next steps:

1. Conduct additional factual investigation
2. Obtain expert opinions to strengthen the grounds
3. Research recent comparable case law
4. Review all trial transcripts for additional grounds

With additional work, these grounds could be strengthened to support a viable appeal."""
        else:
            recommendations = """The current analysis shows primarily weak grounds of appeal. 
It is recommended to:

1. Seek additional evidence or explore alternative legal strategies
2. Consider a comprehensive review of the entire trial record
3. Engage specialist appellate counsel for second opinion
4. Explore other post-conviction remedies if available

Before proceeding with a formal appeal, further groundwork is necessary."""
        
        # Create barrister report
        report = BarristerReport(
            case_id=case_id,
            report_title=report_title,
            executive_summary=executive_summary,
            grounds_identified=grounds_identified,
            legal_analysis_summary=legal_analysis_summary,
            recommendations=recommendations,
            generated_by='AI Legal Analysis System'
        )
        
        db.session.add(report)
        db.session.commit()
        
        logger.info(f"Generated barrister report {report.id} for case {case_id}")
        return jsonify(barrister_report_schema.dump(report)), 201
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error creating barrister report: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating barrister report: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/cases/<int:case_id>/reports', methods=['GET'])
def get_case_reports(case_id):
    """
    Get all barrister reports for a case.
    
    Args:
        case_id: The ID of the case
        
    Returns:
        JSON array of barrister reports
    """
    try:
        case = Case.query.get(case_id)
        
        if not case:
            return jsonify({'error': 'Case not found'}), 404
        
        reports = BarristerReport.query.filter_by(case_id=case_id).order_by(desc(BarristerReport.generated_at)).all()
        
        logger.info(f"Retrieved {len(reports)} barrister reports for case {case_id}")
        return jsonify(barrister_reports_schema.dump(reports)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving barrister reports: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving barrister reports: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/reports/<int:report_id>', methods=['GET'])
def get_barrister_report(report_id):
    """
    Get a specific barrister report by ID.
    
    Args:
        report_id: The ID of the barrister report
        
    Returns:
        JSON response with barrister report data
    """
    try:
        report = BarristerReport.query.get(report_id)
        
        if not report:
            return jsonify({'error': 'Barrister report not found'}), 404
        
        logger.info(f"Retrieved barrister report {report_id}")
        return jsonify(barrister_report_schema.dump(report)), 200
        
    except SQLAlchemyError as e:
        logger.error(f"Database error retrieving barrister report {report_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        logger.error(f"Error retrieving barrister report {report_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


@app.route('/api/reports/<int:report_id>', methods=['DELETE'])
def delete_barrister_report(report_id):
    """
    Delete a barrister report.
    
    Args:
        report_id: The ID of the barrister report to delete
        
    Returns:
        JSON response confirming deletion
    """
    try:
        report = BarristerReport.query.get(report_id)
        
        if not report:
            return jsonify({'error': 'Barrister report not found'}), 404
        
        db.session.delete(report)
        db.session.commit()
        
        logger.info(f"Deleted barrister report {report_id}")
        return jsonify({'message': 'Barrister report deleted successfully'}), 200
        
    except SQLAlchemyError as e:
        db.session.rollback()
        logger.error(f"Database error deleting barrister report {report_id}: {str(e)}")
        return jsonify({'error': 'Database error', 'details': str(e)}), 500
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error deleting barrister report {report_id}: {str(e)}")
        return jsonify({'error': 'Internal server error', 'details': str(e)}), 500


# =====================================================================
# HEALTH CHECK ENDPOINT
# =====================================================================

@app.route('/api/health', methods=['GET'])
def health_check():
    """
    Health check endpoint to verify API is running.
    
    Returns:
        JSON response with health status
    """
    try:
        # Verify database connection by checking if it's usable
        # This is safer than using db.text() which could introduce SQL injection if pattern is copied
        db.session.query(Case).limit(1).all()
        db_status = 'healthy'
    except Exception as e:
        logger.error(f"Database health check failed: {str(e)}")
        db_status = 'unhealthy'
    
    return jsonify({
        'status': 'healthy' if db_status == 'healthy' else 'degraded',
        'database': db_status,
        'timestamp': datetime.utcnow().isoformat()
    }), 200 if db_status == 'healthy' else 503


# =====================================================================
# ERROR HANDLERS
# =====================================================================

@app.errorhandler(404)
def not_found(error):
    """Handle 404 errors."""
    return jsonify({'error': 'Resource not found', 'message': str(error)}), 404


@app.errorhandler(500)
def internal_error(error):
    """Handle 500 errors."""
    db.session.rollback()
    logger.error(f"Internal server error: {str(error)}")
    return jsonify({'error': 'Internal server error', 'message': str(error)}), 500


@app.errorhandler(413)
def request_entity_too_large(error):
    """Handle file too large errors."""
    return jsonify({'error': 'File too large', 'message': 'Maximum file size is 50MB'}), 413


# =====================================================================
# DATABASE INITIALIZATION
# =====================================================================

def init_db():
    """
    Initialize the database by creating all tables.
    """
    with app.app_context():
        try:
            db.create_all()
            logger.info("Database tables created successfully")
        except Exception as e:
            logger.error(f"Error creating database tables: {str(e)}")
            raise


# =====================================================================
# MAIN
# =====================================================================

if __name__ == '__main__':
    # Initialize database
    init_db()
    
    # Get configuration from environment
    host = os.environ.get('FLASK_HOST', '0.0.0.0')
    port = int(os.environ.get('FLASK_PORT', 5000))
    debug = os.environ.get('FLASK_DEBUG', 'False').lower() == 'true'
    
    logger.info(f"Starting Criminal Appeal Case Management System on {host}:{port}")
    
    # Run the application
    app.run(host=host, port=port, debug=debug)
